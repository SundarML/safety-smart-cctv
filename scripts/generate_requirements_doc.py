"""Generate a Word document summarising all hardware and software requirements
for the CCTV PPE Safety Monitoring System."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_table_style(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

def add_header_row(table, headers, bg="2E5FA3"):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)

def add_section_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # === TITLE ===
    title = doc.add_heading("CCTV PPE Safety Monitoring System", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Hardware & Software Requirements")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    date_p = doc.add_paragraph(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # =========================================================
    # 1. PROJECT OVERVIEW
    # =========================================================
    add_section_heading(doc, "1. Project Overview")
    add_body(doc,
        "This system monitors workers on a shopfloor or construction site using CCTV cameras "
        "and AI-based computer vision to detect whether Personal Protective Equipment (PPE) is "
        "being worn. Violations are flagged in real-time with on-screen overlays and logged "
        "automatically with timestamped snapshots."
    )
    add_bullet(doc, "Hardhat / NO-Hardhat", "Detects: ")
    add_bullet(doc, "Safety Vest / NO-Safety Vest")
    add_bullet(doc, "Face Mask / NO-Mask")
    add_bullet(doc, "Gloves, Safety Cones, Persons, Vehicles, Machinery")
    doc.add_paragraph()

    # =========================================================
    # 2. SYSTEM ARCHITECTURE
    # =========================================================
    add_section_heading(doc, "2. System Architecture")
    add_body(doc,
        "The system runs fully offline (no internet or cloud required after initial setup). "
        "All AI inference is performed locally on the edge device."
    )

    arch_lines = [
        "IP Cameras  →  Jetson Orin NX (AI inference)  →  LAN Network Switch",
        "Desktop PC  →  SSH remote access into Jetson  →  View logs & configure",
    ]
    for line in arch_lines:
        add_bullet(doc, line)
    doc.add_paragraph()

    # =========================================================
    # 3. RECOMMENDED HARDWARE
    # =========================================================
    add_section_heading(doc, "3. Recommended Hardware")

    # --- 3.1 Edge AI Device ---
    add_section_heading(doc, "3.1  Edge AI Device — NVIDIA Jetson Orin NX", level=2)
    add_body(doc,
        "The Jetson Orin NX is the recommended compute platform. It is purpose-built for "
        "on-site AI inference, runs 24/7 at low power, and handles multiple camera streams simultaneously."
    )

    tbl = doc.add_table(rows=1, cols=3)
    set_table_style(tbl)
    add_header_row(tbl, ["Specification", "Detail", "Notes"])
    specs = [
        ("Model",         "Jetson Orin NX 16GB",          "16 GB unified memory"),
        ("GPU",           "1024-core NVIDIA Ampere GPU",   "With 32 Tensor Cores"),
        ("AI Performance","100 TOPS",                      "Tera Operations Per Second"),
        ("CPU",           "8-core Arm Cortex-A78AE",       "2 MB L2 + 4 MB L3 cache"),
        ("Power",         "10 – 25 W",                     "Suitable for 24/7 operation"),
        ("Cameras",       "4 – 8 simultaneous streams",    "At 30 FPS per camera"),
        ("Carrier Board", "Seeed Studio reComputer J4012", "Comes bundled with module"),
        ("Estimated Cost","~$600 (module + carrier board)","Single purchase, no assembly"),
    ]
    for i, (spec, detail, note) in enumerate(specs):
        row = tbl.add_row()
        row.cells[0].text = spec
        row.cells[1].text = detail
        row.cells[2].text = note
        row.cells[0].paragraphs[0].runs[0].bold = True
        if i % 2 == 0:
            shade_row(row, "EBF0FB")
    doc.add_paragraph()

    # --- 3.2 Storage ---
    add_section_heading(doc, "3.2  Storage", level=2)
    tbl2 = doc.add_table(rows=1, cols=3)
    set_table_style(tbl2)
    add_header_row(tbl2, ["Item", "Specification", "Estimated Cost"])
    storage = [
        ("NVMe SSD (OS + model)", "128 GB minimum, 256 GB recommended", "~$30"),
        ("External HDD (logs)",   "1 TB (for violation snapshots & CSV logs)", "~$50"),
    ]
    for i, row_data in enumerate(storage):
        row = tbl2.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        if i % 2 == 0:
            shade_row(row, "EBF0FB")
    doc.add_paragraph()

    # --- 3.3 IP Cameras ---
    add_section_heading(doc, "3.3  IP Cameras", level=2)
    add_body(doc,
        "Use PoE (Power over Ethernet) cameras — a single ethernet cable carries both "
        "data and power, eliminating the need for a separate power supply per camera."
    )
    tbl3 = doc.add_table(rows=1, cols=3)
    set_table_style(tbl3)
    add_header_row(tbl3, ["Specification", "Recommended Value", "Notes"])
    cam_specs = [
        ("Resolution",    "1080p (2MP)",             "Sufficient for PPE detection"),
        ("Protocol",      "RTSP",                    "Supported by all major brands"),
        ("Type",          "PoE (802.3af/at)",         "One cable = data + power"),
        ("Brand",         "Hikvision or Dahua",       "Most reliable RTSP support"),
        ("Lens",          "2.8 mm – 4 mm",            "Wide angle for indoor coverage"),
        ("IR Night Vision","Yes (30 m range)",        "For low-light / night shifts"),
        ("Estimated Cost","~$40 – $60 per camera",   ""),
    ]
    for i, (spec, val, note) in enumerate(cam_specs):
        row = tbl3.add_row()
        row.cells[0].text = spec
        row.cells[1].text = val
        row.cells[2].text = note
        row.cells[0].paragraphs[0].runs[0].bold = True
        if i % 2 == 0:
            shade_row(row, "EBF0FB")
    doc.add_paragraph()

    # --- 3.4 Network ---
    add_section_heading(doc, "3.4  Network Equipment", level=2)
    tbl4 = doc.add_table(rows=1, cols=3)
    set_table_style(tbl4)
    add_header_row(tbl4, ["Item", "Specification", "Estimated Cost"])
    net = [
        ("PoE Network Switch", "8-port Gigabit PoE (802.3af/at)", "~$50 – $80"),
        ("Ethernet Cables",    "Cat6, lengths as required",        "~$20"),
        ("Router",             "Existing site router (if any)",    "—"),
    ]
    for i, row_data in enumerate(net):
        row = tbl4.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        if i % 2 == 0:
            shade_row(row, "EBF0FB")
    doc.add_paragraph()

    # --- 3.5 Desktop PC ---
    add_section_heading(doc, "3.5  Desktop / Laptop PC (Management Only)", level=2)
    add_body(doc,
        "The desktop PC is used only for remote access to the Jetson via SSH — "
        "it does not run the AI. Any modern PC or laptop is sufficient."
    )
    tbl5 = doc.add_table(rows=1, cols=2)
    set_table_style(tbl5)
    add_header_row(tbl5, ["Requirement", "Minimum Spec"])
    pc_specs = [
        ("OS",  "Windows 10/11, Ubuntu, or macOS"),
        ("RAM", "8 GB"),
        ("Network", "Connected to the same LAN as Jetson"),
        ("Software", "SSH client (built into Windows 10+, PowerShell)"),
    ]
    for i, (req, val) in enumerate(pc_specs):
        row = tbl5.add_row()
        row.cells[0].text = req
        row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].bold = True
        if i % 2 == 0:
            shade_row(row, "EBF0FB")
    doc.add_paragraph()

    # =========================================================
    # 4. COMPLETE SHOPPING LIST
    # =========================================================
    add_section_heading(doc, "4. Complete Shopping List (4-Camera Setup)")
    tbl6 = doc.add_table(rows=1, cols=4)
    set_table_style(tbl6)
    add_header_row(tbl6, ["Item", "Model / Spec", "Qty", "Est. Cost (USD)"])
    shopping = [
        ("Jetson Orin NX 16GB + Carrier Board", "Seeed reComputer J4012",         "1", "$600"),
        ("NVMe SSD",                             "128 GB NVMe M.2",                "1", "$30"),
        ("IP Camera (PoE, 1080p)",               "Hikvision DS-2CD2143G2-I",       "4", "$160"),
        ("PoE Gigabit Switch",                   "TP-Link TL-SG1008PE (8-port)",   "1", "$65"),
        ("Cat6 Ethernet Cable",                  "Various lengths",                "—", "$20"),
        ("External HDD (logs)",                  "1 TB USB HDD",                   "1", "$50"),
        ("",                                     "",                               "",  ""),
        ("TOTAL",                                "",                               "",  "~$925"),
    ]
    for i, (item, model, qty, cost) in enumerate(shopping):
        row = tbl6.add_row()
        row.cells[0].text = item
        row.cells[1].text = model
        row.cells[2].text = qty
        row.cells[3].text = cost
        if item == "TOTAL":
            for cell in row.cells:
                cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else None
                run = cell.paragraphs[0].runs
                if run:
                    run[0].bold = True
            shade_row(row, "D9E1F2")
        elif i % 2 == 0 and item:
            shade_row(row, "EBF0FB")
    doc.add_paragraph()

    # =========================================================
    # 5. GPU REQUIREMENTS (if using a laptop instead)
    # =========================================================
    add_section_heading(doc, "5. Alternative — GPU Laptop Instead of Jetson")
    add_body(doc,
        "If portability is preferred over dedicated edge deployment, a GPU laptop can run "
        "the system. Minimum GPU spec is RTX 4050 with 6 GB VRAM."
    )
    tbl7 = doc.add_table(rows=1, cols=4)
    set_table_style(tbl7)
    add_header_row(tbl7, ["GPU", "VRAM", "Cameras Supported", "Verdict"])
    gpus = [
        ("RTX 3050", "4 GB", "2 – 3", "Minimum — not recommended"),
        ("RTX 4050", "6 GB", "3 – 4", "Entry level — acceptable"),
        ("RTX 4060", "8 GB", "6 – 8", "Sweet spot — recommended"),
        ("RTX 4070", "12 GB","10+",   "Overkill unless scaling up"),
    ]
    for i, row_data in enumerate(gpus):
        row = tbl7.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = val
        if i % 2 == 0:
            shade_row(row, "EBF0FB")
    add_body(doc, "\nRecommended laptop model: ASUS TUF Gaming A15 / Lenovo Legion 5 with RTX 4060 (~$900–$1,100).")
    doc.add_paragraph()

    # =========================================================
    # 6. SOFTWARE REQUIREMENTS
    # =========================================================
    add_section_heading(doc, "6. Software Requirements")
    tbl8 = doc.add_table(rows=1, cols=3)
    set_table_style(tbl8)
    add_header_row(tbl8, ["Component", "Software / Version", "Notes"])
    software = [
        ("Operating System",  "JetPack 6.x (Ubuntu 22.04-based)", "NVIDIA official OS for Jetson"),
        ("Python",            "Python 3.10+",                      "Pre-installed on JetPack"),
        ("AI Framework",      "Ultralytics YOLOv8",                "pip install ultralytics"),
        ("Computer Vision",   "OpenCV 4.8+",                       "pip install opencv-python"),
        ("Dataset Platform",  "Roboflow",                          "For dataset download only"),
        ("PPE Model",         "YOLOv8n fine-tuned on PPE data",    "Trained and saved as ppe.pt"),
        ("Config Format",     "YAML (settings.yaml)",              "Cameras, thresholds, classes"),
        ("Log Format",        "CSV + JPEG snapshots",              "Stored in logs/ directory"),
    ]
    for i, (comp, sw, note) in enumerate(software):
        row = tbl8.add_row()
        row.cells[0].text = comp
        row.cells[1].text = sw
        row.cells[2].text = note
        row.cells[0].paragraphs[0].runs[0].bold = True
        if i % 2 == 0:
            shade_row(row, "EBF0FB")
    doc.add_paragraph()

    # =========================================================
    # 7. NETWORK LAYOUT
    # =========================================================
    add_section_heading(doc, "7. Site Network Layout")
    add_body(doc, "All devices are connected to the same local area network (LAN):")
    network_layout = [
        "Internet Router  →  connects site to internet (optional for this system)",
        "PoE Network Switch  →  central hub connecting all devices",
        "Jetson Orin NX  →  192.168.1.50  (runs the AI monitoring software)",
        "IP Camera 1      →  192.168.1.101",
        "IP Camera 2      →  192.168.1.102",
        "IP Camera 3      →  192.168.1.103",
        "IP Camera 4      →  192.168.1.104",
        "Desktop PC       →  192.168.1.10  (remote management via SSH)",
    ]
    for line in network_layout:
        add_bullet(doc, line)
    doc.add_paragraph()

    # =========================================================
    # 8. HOW IT WORKS (SUMMARY)
    # =========================================================
    add_section_heading(doc, "8. How It Works — Summary")
    steps = [
        ("Camera streams video", "IP cameras stream 1080p RTSP video continuously over LAN."),
        ("Jetson captures frames", "The monitoring software reads frames from each camera in a background thread."),
        ("YOLOv8 runs inference", "The trained PPE model (ppe.pt) processes each frame and outputs bounding boxes with class labels."),
        ("Violation check", "Any detection matching NO-Hardhat, NO-Safety Vest, or NO-Mask is flagged as a violation."),
        ("Display overlay", "Red boxes = violations, Green boxes = compliant PPE. Status bar shows ALL CLEAR or VIOLATION."),
        ("Logging", "Each violation is written to violations.csv with timestamp, camera name, and violation type. A snapshot image is saved."),
        ("Remote access", "Operators can SSH into the Jetson from any PC on the network to check logs, restart the service, or update the model."),
    ]
    for i, (title_text, desc) in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        run_title = p.add_run(f"{title_text}: ")
        run_title.bold = True
        p.add_run(desc)
    doc.add_paragraph()

    # =========================================================
    # 9. KEY ADVANTAGES
    # =========================================================
    add_section_heading(doc, "9. Key Advantages")
    advantages = [
        ("Fully offline", "No internet, no cloud subscription, no recurring fees after purchase."),
        ("Real-time", "Detects violations frame by frame as they happen."),
        ("Scalable", "Add more cameras by updating the config file — no code changes needed."),
        ("Low power", "Jetson Orin NX uses 10–25W, suitable for 24/7 operation."),
        ("Automatic logging", "All violations are timestamped and saved without manual intervention."),
        ("Open source stack", "Built on YOLOv8, OpenCV, Python — no proprietary software licenses."),
    ]
    for bold, desc in advantages:
        add_bullet(doc, f" {desc}", bold_prefix=bold + ":")
    doc.add_paragraph()

    # --- Footer note ---
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Note: All prices are approximate USD estimates as of 2026 and may vary by region and supplier. "
        "The system runs on Python 3.10+ and is compatible with JetPack 6.x on Jetson Orin NX."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Save
    output_path = "CCTV_PPE_System_Requirements.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")


if __name__ == "__main__":
    build()
